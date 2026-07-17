"""從 Word .docx 抽取文字（段落與表格），預設轉為 Markdown。

**預設輸出 Markdown（.md）**：以 mammoth 轉出、保留標題層級，供後續 Read 閱讀。
需純文字（段落＋表格、以 python-docx 抽取）時用 `--txt`。

輸出路徑：省略 `-o` 時，由輸入檔名主幹加副檔名，落在環境變數 `SDE_OUT_DIR`
指定的資料夾（預設 `extracted`）；`-o` 顯式指定時以其為準。

需在裝有 python-docx 與 mammoth 的 conda 環境下執行（預設環境名 `research`）：
    conda run -n research pip install python-docx mammoth

與 extract_pdf.py 相同理由：一律將結果寫入 UTF-8 檔案再由其他工具讀取，
避免 Windows cp950 編碼錯誤與 `conda run -c` 多行限制。

用法範例：
    # 以 mammoth 轉為 Markdown（預設、保留標題層級；輸出 extracted/draft.md）
    conda run -n research python scripts/extract_docx.py draft.docx

    # 以 python-docx 抽取段落＋表格純文字（除錯或需保留表格原貌時）
    conda run -n research python scripts/extract_docx.py draft.docx --txt
"""
import argparse
import os
import sys


def extract_with_python_docx(path: str) -> str:
    """逐段落輸出；表格以 tab 分隔每列，置於 [TABLE] 標記之間。"""
    import docx  # python-docx

    doc = docx.Document(path)
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # 保留標題層級提示，方便辨識章節
            style = (para.style.name or "") if para.style else ""
            prefix = f"[{style}] " if style.lower().startswith("heading") else ""
            lines.append(prefix + text)
    for ti, table in enumerate(doc.tables, 1):
        lines.append(f"\n[TABLE {ti}]")
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            lines.append("\t".join(cells))
        lines.append("[/TABLE]")
    return "\n".join(lines)


def extract_with_mammoth(path: str) -> str:
    """轉為 Markdown，保留標題與清單結構。"""
    import mammoth

    with open(path, "rb") as f:
        result = mammoth.convert_to_markdown(f)
    return result.value


def default_out(src: str, ext: str) -> str:
    """省略 -o 時的輸出路徑：SDE_OUT_DIR（預設 extracted）/ 輸入主幹 + ext。"""
    out_dir = os.environ.get("SDE_OUT_DIR", "extracted")
    stem = os.path.splitext(os.path.basename(src))[0]
    return os.path.join(out_dir, stem + ext)


def main() -> None:
    ap = argparse.ArgumentParser(description="抽取 .docx 文字（預設輸出 Markdown 檔）")
    ap.add_argument("docx", help=".docx 路徑，例如 draft.docx")
    ap.add_argument("--txt", action="store_true",
                    help="改用 python-docx 抽段落＋表格純文字（預設用 mammoth 轉 Markdown）")
    ap.add_argument("-o", "--out",
                    help="輸出檔路徑（UTF-8）；省略則用 SDE_OUT_DIR（預設 extracted）")
    args = ap.parse_args()

    text = (extract_with_python_docx(args.docx) if args.txt
            else extract_with_mammoth(args.docx))

    out = args.out if args.out else default_out(args.docx, ".txt" if args.txt else ".md")
    out_dir = os.path.dirname(out)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    with open(out, "w", encoding="utf-8") as f:
        f.write(text)

    safe = out.encode("ascii", "replace").decode("ascii")
    print(f"done -> {safe}", file=sys.stderr)


if __name__ == "__main__":
    main()
